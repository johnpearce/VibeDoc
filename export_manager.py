"""
VibeDoc å¤šæ ¼å¼å¯¼å‡ºç®¡ç†å™¨
æ”¯æŒ Ma# PDF å¯¼å‡º
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# é«˜çº§PDFå¯¼å‡º - ç§»é™¤weasyprintä¾èµ–ï¼Œä½¿ç”¨reportlab
WEASYPRINT_AVAILABLE = FalseF æ ¼å¼çš„æ–‡æ¡£å¯¼å‡º
"""

import os
import io
import re
import zipfile
import tempfile
from datetime import datetime
from typing import Dict, Tuple, Optional, Any
import logging

# æ ¸å¿ƒä¾èµ–
import markdown
import html2text

# Word å¯¼å‡º
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF å¯¼å‡º
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# é«˜çº§PDFå¯¼å‡ºï¼ˆå¤‡ç”¨æ–¹æ¡ˆï¼‰ - ç§»é™¤weasyprintä¾èµ–
WEASYPRINT_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportManager:
    """å¤šæ ¼å¼å¯¼å‡ºç®¡ç†å™¨"""
    
    def __init__(self):
        self.supported_formats = ['markdown', 'html']
        
        # æ£€æŸ¥å¯é€‰ä¾èµ–
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')
        if PDF_AVAILABLE:
            self.supported_formats.append('pdf')
            
        logger.info(f"ğŸ“„ ExportManager åˆå§‹åŒ–å®Œæˆï¼Œæ”¯æŒæ ¼å¼: {', '.join(self.supported_formats)}")
    
    def get_supported_formats(self) -> list:
        """è·å–æ”¯æŒçš„å¯¼å‡ºæ ¼å¼"""
        return self.supported_formats.copy()
    
    def export_to_markdown(self, content: str, metadata: Optional[Dict] = None) -> str:
        """
        å¯¼å‡ºä¸º Markdown æ ¼å¼ï¼ˆæ¸…ç†å’Œä¼˜åŒ–ï¼‰
        
        Args:
            content: åŸå§‹å†…å®¹
            metadata: å…ƒæ•°æ®ä¿¡æ¯
            
        Returns:
            str: ä¼˜åŒ–åçš„ Markdown å†…å®¹
        """
        try:
            # æ·»åŠ æ–‡æ¡£å¤´éƒ¨ä¿¡æ¯
            if metadata:
                header = f"""---
title: {metadata.get('title', 'VibeDoc å¼€å‘è®¡åˆ’')}
author: {metadata.get('author', 'VibeDoc AI Agent')}
date: {metadata.get('date', datetime.now().strftime('%Y-%m-%d'))}
generator: VibeDoc AI Agent v1.0
---

"""
                content = header + content
            
            # æ¸…ç†å’Œä¼˜åŒ–å†…å®¹
            content = self._clean_markdown_content(content)
            
            logger.info("âœ… Markdown å¯¼å‡ºæˆåŠŸ")
            return content
            
        except Exception as e:
            logger.error(f"âŒ Markdown å¯¼å‡ºå¤±è´¥: {e}")
            return content  # è¿”å›åŸå§‹å†…å®¹
    
    def export_to_html(self, content: str, metadata: Optional[Dict] = None) -> str:
        """
        å¯¼å‡ºä¸º HTML æ ¼å¼ï¼ˆå¸¦æ ·å¼ï¼‰
        
        Args:
            content: Markdown å†…å®¹
            metadata: å…ƒæ•°æ®ä¿¡æ¯
            
        Returns:
            str: å®Œæ•´çš„ HTML å†…å®¹
        """
        try:
            # é…ç½® Markdown æ‰©å±•
            md = markdown.Markdown(
                extensions=[
                    'markdown.extensions.extra',
                    'markdown.extensions.codehilite',
                    'markdown.extensions.toc',
                    'markdown.extensions.tables'
                ],
                extension_configs={
                    'codehilite': {
                        'css_class': 'highlight',
                        'use_pygments': False
                    },
                    'toc': {
                        'title': 'ç›®å½•'
                    }
                }
            )
            
            # è½¬æ¢ Markdown åˆ° HTML
            html_content = md.convert(content)
            
            # ç”Ÿæˆå®Œæ•´çš„ HTML æ–‡æ¡£
            title = metadata.get('title', 'VibeDoc å¼€å‘è®¡åˆ’') if metadata else 'VibeDoc å¼€å‘è®¡åˆ’'
            author = metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'
            date = metadata.get('date', datetime.now().strftime('%Y-%m-%d')) if metadata else datetime.now().strftime('%Y-%m-%d')
            
            full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <meta name="author" content="{author}">
    <meta name="generator" content="VibeDoc AI Agent">
    <style>
        {self._get_html_styles()}
    </style>
    <!-- Mermaid æ”¯æŒ -->
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10.6.1/dist/mermaid.min.js"></script>
    <script>
        document.addEventListener('DOMContentLoaded', function() {{
            mermaid.initialize({{ 
                startOnLoad: true,
                theme: 'default',
                securityLevel: 'loose',
                flowchart: {{ useMaxWidth: true }}
            }});
        }});
    </script>
</head>
<body>
    <div class="container">
        <header class="document-header">
            <h1>{title}</h1>
            <div class="meta-info">
                <span class="author">ğŸ“ {author}</span>
                <span class="date">ğŸ“… {date}</span>
                <span class="generator">ğŸ¤– Generated by VibeDoc AI Agent</span>
            </div>
        </header>
        
        <main class="content">
            {html_content}
        </main>
        
        <footer class="document-footer">
            <p>æœ¬æ–‡æ¡£ç”± <strong>VibeDoc AI Agent</strong> ç”Ÿæˆ | ç”Ÿæˆæ—¶é—´: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </footer>
    </div>
</body>
</html>"""
            
            logger.info("âœ… HTML å¯¼å‡ºæˆåŠŸ")
            return full_html
            
        except Exception as e:
            logger.error(f"âŒ HTML å¯¼å‡ºå¤±è´¥: {e}")
            # ç®€å•çš„ HTML å¤‡ç”¨æ–¹æ¡ˆ
            return f"""<!DOCTYPE html>
<html><head><title>VibeDoc å¼€å‘è®¡åˆ’</title></head>
<body><pre>{content}</pre></body></html>"""
    
    def export_to_docx(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """
        å¯¼å‡ºä¸º Word æ–‡æ¡£æ ¼å¼
        
        Args:
            content: Markdown å†…å®¹
            metadata: å…ƒæ•°æ®ä¿¡æ¯
            
        Returns:
            bytes: Word æ–‡æ¡£çš„äºŒè¿›åˆ¶æ•°æ®
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx æœªå®‰è£…ï¼Œæ— æ³•å¯¼å‡º Word æ ¼å¼")
        
        try:
            # åˆ›å»ºæ–°æ–‡æ¡£
            doc = Document()
            
            # è®¾ç½®æ–‡æ¡£å±æ€§
            properties = doc.core_properties
            properties.title = metadata.get('title', 'VibeDoc å¼€å‘è®¡åˆ’') if metadata else 'VibeDoc å¼€å‘è®¡åˆ’'
            properties.author = metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'
            properties.subject = 'AIé©±åŠ¨çš„æ™ºèƒ½å¼€å‘è®¡åˆ’'
            properties.comments = 'Generated by VibeDoc AI Agent'
            
            # æ·»åŠ æ ‡é¢˜
            title = doc.add_heading(properties.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # æ·»åŠ å…ƒä¿¡æ¯
            doc.add_paragraph()
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"ğŸ“ ä½œè€…: {properties.author}").bold = True
            meta_para.add_run("\n")
            meta_para.add_run(f"ğŸ“… ç”Ÿæˆæ—¶é—´: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").bold = True
            meta_para.add_run("\n")
            meta_para.add_run("ğŸ¤– ç”Ÿæˆå·¥å…·: VibeDoc AI Agent").bold = True
            
            doc.add_paragraph()
            doc.add_paragraph("â”€" * 50)
            doc.add_paragraph()
            
            # è§£æå’Œæ·»åŠ å†…å®¹
            self._parse_markdown_to_docx(doc, content)
            
            # æ·»åŠ é¡µè„š
            doc.add_paragraph()
            doc.add_paragraph("â”€" * 50)
            footer_para = doc.add_paragraph()
            footer_para.add_run("æœ¬æ–‡æ¡£ç”± VibeDoc AI Agent è‡ªåŠ¨ç”Ÿæˆ").italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # ä¿å­˜åˆ°å†…å­˜
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            
            logger.info("âœ… Word æ–‡æ¡£å¯¼å‡ºæˆåŠŸ")
            return doc_stream.getvalue()
            
        except Exception as e:
            logger.error(f"âŒ Word å¯¼å‡ºå¤±è´¥: {e}")
            raise
    
    def export_to_pdf(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """
        å¯¼å‡ºä¸º PDF æ ¼å¼
        
        Args:
            content: Markdown å†…å®¹  
            metadata: å…ƒæ•°æ®ä¿¡æ¯
            
        Returns:
            bytes: PDF æ–‡æ¡£çš„äºŒè¿›åˆ¶æ•°æ®
        """
        if PDF_AVAILABLE:
            return self._export_pdf_reportlab(content, metadata)
        else:
            raise ImportError("PDF å¯¼å‡ºä¾èµ–æœªå®‰è£…")
    
    def create_multi_format_export(self, content: str, formats: list = None, metadata: Optional[Dict] = None) -> bytes:
        """
        åˆ›å»ºå¤šæ ¼å¼å¯¼å‡ºçš„ ZIP åŒ…
        
        Args:
            content: åŸå§‹å†…å®¹
            formats: è¦å¯¼å‡ºçš„æ ¼å¼åˆ—è¡¨ï¼Œé»˜è®¤ä¸ºæ‰€æœ‰æ”¯æŒçš„æ ¼å¼
            metadata: å…ƒæ•°æ®ä¿¡æ¯
            
        Returns:
            bytes: ZIP æ–‡ä»¶çš„äºŒè¿›åˆ¶æ•°æ®
        """
        if formats is None:
            formats = self.supported_formats
        
        # éªŒè¯æ ¼å¼
        invalid_formats = set(formats) - set(self.supported_formats)
        if invalid_formats:
            raise ValueError(f"ä¸æ”¯æŒçš„æ ¼å¼: {', '.join(invalid_formats)}")
        
        try:
            # åˆ›å»ºå†…å­˜ä¸­çš„ ZIP æ–‡ä»¶
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # ç”ŸæˆåŸºç¡€æ–‡ä»¶å
                base_name = metadata.get('title', 'vibedoc_plan') if metadata else 'vibedoc_plan'
                base_name = re.sub(r'[^\w\-_\.]', '_', base_name)  # æ¸…ç†æ–‡ä»¶å
                
                # å¯¼å‡ºå„ç§æ ¼å¼
                for fmt in formats:
                    try:
                        if fmt == 'markdown':
                            file_content = self.export_to_markdown(content, metadata)
                            zip_file.writestr(f"{base_name}.md", file_content.encode('utf-8'))
                            
                        elif fmt == 'html':
                            file_content = self.export_to_html(content, metadata)
                            zip_file.writestr(f"{base_name}.html", file_content.encode('utf-8'))
                            
                        elif fmt == 'docx' and DOCX_AVAILABLE:
                            file_content = self.export_to_docx(content, metadata)
                            zip_file.writestr(f"{base_name}.docx", file_content)
                            
                        elif fmt == 'pdf' and PDF_AVAILABLE:
                            file_content = self.export_to_pdf(content, metadata)
                            zip_file.writestr(f"{base_name}.pdf", file_content)
                            
                    except Exception as e:
                        logger.warning(f"âš ï¸ æ ¼å¼ {fmt} å¯¼å‡ºå¤±è´¥: {e}")
                        # åœ¨ ZIP ä¸­æ·»åŠ é”™è¯¯ä¿¡æ¯æ–‡ä»¶
                        error_msg = f"æ ¼å¼ {fmt} å¯¼å‡ºå¤±è´¥:\n{str(e)}\n\nè¯·æ£€æŸ¥ç›¸å…³ä¾èµ–æ˜¯å¦æ­£ç¡®å®‰è£…ã€‚"
                        zip_file.writestr(f"ERROR_{fmt}.txt", error_msg.encode('utf-8'))
                
                # æ·»åŠ è¯´æ˜æ–‡ä»¶
                readme_content = f"""# VibeDoc å¯¼å‡ºæ–‡ä»¶åŒ…

## ğŸ“‹ æ–‡ä»¶è¯´æ˜
æœ¬å‹ç¼©åŒ…åŒ…å«äº†æ‚¨çš„å¼€å‘è®¡åˆ’çš„å¤šç§æ ¼å¼å¯¼å‡ºï¼š

### ğŸ“„ æ”¯æŒçš„æ ¼å¼ï¼š
- **Markdown (.md)**: åŸå§‹æ ¼å¼ï¼Œæ”¯æŒæ‰€æœ‰ Markdown è¯­æ³•
- **HTML (.html)**: ç½‘é¡µæ ¼å¼ï¼ŒåŒ…å«æ ·å¼å’Œ Mermaid å›¾è¡¨æ”¯æŒ
- **Word (.docx)**: Microsoft Word æ–‡æ¡£æ ¼å¼
- **PDF (.pdf)**: ä¾¿æºå¼æ–‡æ¡£æ ¼å¼

### ğŸ¤– ç”Ÿæˆä¿¡æ¯ï¼š
- ç”Ÿæˆæ—¶é—´: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- ç”Ÿæˆå·¥å…·: VibeDoc AI Agent v1.0
- é¡¹ç›®åœ°å€: https://github.com/JasonRobertDestiny/VibeDocs

### ğŸ’¡ ä½¿ç”¨å»ºè®®ï¼š
1. ä¼˜å…ˆä½¿ç”¨ HTML æ ¼å¼æŸ¥çœ‹ï¼Œæ”¯æŒæœ€ä½³çš„è§†è§‰æ•ˆæœ
2. ä½¿ç”¨ Markdown æ ¼å¼è¿›è¡Œè¿›ä¸€æ­¥ç¼–è¾‘
3. ä½¿ç”¨ Word æ ¼å¼è¿›è¡Œæ­£å¼æ–‡æ¡£å¤„ç†
4. ä½¿ç”¨ PDF æ ¼å¼è¿›è¡Œåˆ†äº«å’Œæ‰“å°

---
æ„Ÿè°¢ä½¿ç”¨ VibeDoc AI Agentï¼
"""
                zip_file.writestr("README.md", readme_content.encode('utf-8'))
            
            zip_buffer.seek(0)
            logger.info(f"âœ… å¤šæ ¼å¼å¯¼å‡ºæˆåŠŸï¼ŒåŒ…å« {len(formats)} ç§æ ¼å¼")
            return zip_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"âŒ å¤šæ ¼å¼å¯¼å‡ºå¤±è´¥: {e}")
            raise
    
    def _clean_markdown_content(self, content: str) -> str:
        """æ¸…ç†å’Œä¼˜åŒ– Markdown å†…å®¹"""
        # ä¿®å¤å¸¸è§çš„æ ¼å¼é—®é¢˜
        content = re.sub(r'\n{3,}', '\n\n', content)  # ç§»é™¤å¤šä½™ç©ºè¡Œ
        content = re.sub(r'(?m)^[ \t]+$', '', content)  # ç§»é™¤åªæœ‰ç©ºæ ¼çš„è¡Œ
        content = content.strip()
        
        return content
    
    def _get_html_styles(self) -> str:
        """è·å– HTML æ ·å¼"""
        return """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Hiragino Sans GB', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f8fafc;
        }
        
        .container {
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
            border-radius: 8px;
            margin-top: 20px;
            margin-bottom: 20px;
        }
        
        .document-header {
            text-align: center;
            border-bottom: 3px solid #667eea;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        
        .document-header h1 {
            color: #667eea;
            font-size: 2.2em;
            margin-bottom: 15px;
        }
        
        .meta-info {
            display: flex;
            justify-content: center;
            gap: 20px;
            flex-wrap: wrap;
            color: #666;
            font-size: 0.9em;
        }
        
        .content h1, .content h2, .content h3, .content h4 {
            color: #2d3748;
            margin-top: 2em;
            margin-bottom: 1em;
        }
        
        .content h1 { border-bottom: 2px solid #667eea; padding-bottom: 0.5em; }
        .content h2 { border-bottom: 1px solid #e2e8f0; padding-bottom: 0.3em; }
        
        .content p {
            margin-bottom: 1em;
            text-align: justify;
        }
        
        .content ul, .content ol {
            margin-bottom: 1em;
            padding-left: 2em;
        }
        
        .content li {
            margin-bottom: 0.5em;
        }
        
        .content pre {
            background: #2d3748;
            color: #e2e8f0;
            padding: 1em;
            border-radius: 6px;
            overflow-x: auto;
            margin: 1em 0;
        }
        
        .content code {
            background: #f7fafc;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: 'SFMono-Regular', Consolas, monospace;
        }
        
        .content table {
            width: 100%;
            border-collapse: collapse;
            margin: 1em 0;
        }
        
        .content th, .content td {
            border: 1px solid #e2e8f0;
            padding: 0.75em;
            text-align: left;
        }
        
        .content th {
            background: #f7fafc;
            font-weight: 600;
        }
        
        .content blockquote {
            border-left: 4px solid #667eea;
            margin: 1em 0;
            padding-left: 1em;
            color: #666;
            font-style: italic;
        }
        
        .mermaid {
            text-align: center;
            margin: 2em 0;
        }
        
        .document-footer {
            margin-top: 3em;
            padding-top: 20px;
            border-top: 1px solid #e2e8f0;
            text-align: center;
            color: #666;
            font-size: 0.9em;
        }
        
        @media (max-width: 768px) {
            .container {
                margin: 10px;
                padding: 15px;
            }
            
            .meta-info {
                flex-direction: column;
                gap: 10px;
            }
        }
        """
    
    def _parse_markdown_to_docx(self, doc: "Document", content: str):
        """è§£æ Markdown å†…å®¹å¹¶æ·»åŠ åˆ° Word æ–‡æ¡£"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
                
            # æ ‡é¢˜å¤„ç†
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()
                if level <= 6:
                    doc.add_heading(title_text, level)
                    continue
            
            # ä»£ç å—å¤„ç†ï¼ˆç®€åŒ–ï¼‰
            if line.startswith('```'):
                continue
                
            # åˆ—è¡¨å¤„ç†
            if line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                para = doc.add_paragraph(text, style='List Bullet')
                continue
                
            if re.match(r'^\d+\.', line):
                text = re.sub(r'^\d+\.\s*', '', line)
                para = doc.add_paragraph(text, style='List Number')
                continue
            
            # æ™®é€šæ®µè½
            if line:
                # ç®€å•çš„ç²—ä½“å’Œæ–œä½“å¤„ç†
                line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # ç§»é™¤ç²—ä½“æ ‡è®°ï¼ŒWord ä¸­åç»­å¯ä»¥æ‰‹åŠ¨è®¾ç½®
                line = re.sub(r'\*(.*?)\*', r'\1', line)      # ç§»é™¤æ–œä½“æ ‡è®°
                doc.add_paragraph(line)
    
    def _export_pdf_reportlab(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """ä½¿ç”¨ ReportLab å¯¼å‡º PDF"""
        try:
            buffer = io.BytesIO()
            
            # åˆ›å»º PDF æ–‡æ¡£
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                topMargin=1*inch,
                bottomMargin=1*inch,
                leftMargin=1*inch,
                rightMargin=1*inch
            )
            
            # æ ·å¼è®¾ç½®
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=20,
                spaceAfter=30,
                alignment=1  # å±…ä¸­
            )
            
            # æ„å»ºå†…å®¹
            story = []
            
            # æ·»åŠ æ ‡é¢˜
            title = metadata.get('title', 'VibeDoc å¼€å‘è®¡åˆ’') if metadata else 'VibeDoc å¼€å‘è®¡åˆ’'
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            # æ·»åŠ å…ƒä¿¡æ¯
            meta_text = f"""
            ä½œè€…: {metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'}<br/>
            ç”Ÿæˆæ—¶é—´: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
            ç”Ÿæˆå·¥å…·: VibeDoc AI Agent
            """
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 30))
            
            # ç®€å•å¤„ç† Markdown å†…å®¹ï¼ˆåŸºç¡€ç‰ˆæœ¬ï¼‰
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 12))
                    continue
                    
                if line.startswith('#'):
                    # æ ‡é¢˜
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()
                    if level == 1:
                        story.append(Paragraph(title_text, styles['Heading1']))
                    elif level == 2:
                        story.append(Paragraph(title_text, styles['Heading2']))
                    else:
                        story.append(Paragraph(title_text, styles['Heading3']))
                else:
                    # æ™®é€šæ®µè½
                    story.append(Paragraph(line, styles['Normal']))
                    
                story.append(Spacer(1, 6))
            
            # ç”Ÿæˆ PDF
            doc.build(story)
            buffer.seek(0)
            
            logger.info("âœ… PDF å¯¼å‡ºæˆåŠŸï¼ˆReportLabï¼‰")
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"âŒ ReportLab PDF å¯¼å‡ºå¤±è´¥: {e}")
            raise

# å…¨å±€å¯¼å‡ºç®¡ç†å™¨å®ä¾‹
export_manager = ExportManager()